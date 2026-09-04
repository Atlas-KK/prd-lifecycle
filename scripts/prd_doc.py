#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PRD 工作文档版本助手 —— 管理版本自增、修订历史，并导出 .docx 交付稿。

工作文档为 Markdown，顶部含"版本历史"表与"当前版本"标记，下接各章节。
本脚本只负责确定性、易出错的部分（版本号自增 + 历史记录 + 导出）；
章节正文由 Claude 在工作流中用编辑工具直接维护。

命令行：
    python3 prd_doc.py init   <prd.md> "<文档标题>" ["修订人"]
    python3 prd_doc.py bump   <prd.md> <minor|major> "<阶段/步骤>" "<修订内容>" ["修订人"]
    python3 prd_doc.py export <prd.md> <out.docx>
    python3 prd_doc.py version <prd.md>          # 打印当前版本
"""
import re, sys, datetime, subprocess, shutil

HEADER = "| 版本 | 日期 | 阶段/步骤 | 修订内容 | 修订人 |"
SEP = "|---|---|---|---|---|"
SKELETON_SECTIONS = [
    "1. 需求背景与核心问题",
    "2. 用户故事与目标用户",
    "3. 产品价值",
    "4. 功能方案设计",
    "5. 边界条件",
    "6. 验收标准",
]


def today():
    return datetime.date.today().isoformat()


def _read(path):
    with open(path, encoding="utf-8") as f:
        return f.read()


def _write(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def init_doc(path, title, author="产品经理"):
    rows = "\n".join(
        "## %s\n（待补全）\n" % s for s in SKELETON_SECTIONS
    )
    content = (
        "# %s\n\n## 版本历史\n\n%s\n%s\n"
        "| v1.0 | %s | 基线导入 | 导入原始需求文档作为调优补全基线 | %s |\n\n"
        "> **当前版本：v1.0**\n\n---\n\n%s"
    ) % (title, HEADER, SEP, today(), author, rows)
    _write(path, content)
    return "v1.0"


def current_version(path):
    """最新版本 = 历史表中最靠上的数据行。"""
    text = _read(path)
    m = re.search(r"\|\s*v(\d+)\.(\d+)\s*\|", text)
    return (int(m.group(1)), int(m.group(2))) if m else (1, 0)


def bump(path, level, step, summary, author="产品经理"):
    major, minor = current_version(path)
    if level == "major":
        major, minor = major + 1, 0
    else:
        minor += 1
    newver = "v%d.%d" % (major, minor)
    text = _read(path)
    # 在分隔行后插入新行（成为最新一行）
    sep_re = re.compile(r"(\|\s*版本\s*\|.*\n\|[-|: ]+\|\n)")
    newrow = "| %s | %s | %s | %s | %s |\n" % (newver, today(), step, summary, author)
    if not sep_re.search(text):
        raise SystemExit("未找到版本历史表，请先 init")
    text = sep_re.sub(lambda m: m.group(1) + newrow, text, count=1)
    text = re.sub(r">\s*\*\*当前版本：v[\d.]+\*\*",
                  "> **当前版本：%s**" % newver, text)
    _write(path, text)
    return newver


def export_docx(md_path, docx_path):
    if shutil.which("pandoc"):
        subprocess.run(["pandoc", md_path, "-o", docx_path], check=True)
        return docx_path
    # 无 pandoc 时的简易回退（python-docx）
    from docx import Document
    doc = Document()
    in_table = False
    for line in _read(md_path).splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip().startswith("|"):
            # 简易表格：逐行作为段落保留（保证不丢信息）
            doc.add_paragraph(line.strip())
        elif line.strip() and not line.strip().startswith("---"):
            doc.add_paragraph(line)
    doc.save(docx_path)
    return docx_path


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        return
    cmd, path = sys.argv[1], sys.argv[2]
    if cmd == "init":
        title = sys.argv[3] if len(sys.argv) > 3 else "产品需求文档"
        author = sys.argv[4] if len(sys.argv) > 4 else "产品经理"
        print("已建基线:", init_doc(path, title, author), "->", path)
    elif cmd == "bump":
        level = sys.argv[3]
        step = sys.argv[4] if len(sys.argv) > 4 else ""
        summary = sys.argv[5] if len(sys.argv) > 5 else ""
        author = sys.argv[6] if len(sys.argv) > 6 else "产品经理"
        print("已自增至:", bump(path, level, step, summary, author))
    elif cmd == "export":
        out = sys.argv[3] if len(sys.argv) > 3 else "PRD.docx"
        print("已导出:", export_docx(path, out))
    elif cmd == "version":
        mj, mn = current_version(path)
        print("v%d.%d" % (mj, mn))
    else:
        print(__doc__)


if __name__ == "__main__":
    main()
